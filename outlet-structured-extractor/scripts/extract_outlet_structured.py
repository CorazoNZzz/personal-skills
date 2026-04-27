#!/usr/bin/env python3
"""Extract structured emission/facility data from enterprise .docx summary materials.

Trial usage:
    python scripts/extract_outlet_structured.py --limit 5

Full run:
    python scripts/extract_outlet_structured.py --all
"""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


TYPE_ORDER = [
    "document_section",
    "enterprise_profile",
    "section_summary",
    "outlet_overview",
    "issue_identification",
    "reduction_space_identification",
    "organized_outlet",
    "tank",
    "loading",
    "wastewater_surface",
    "fugitive_source",
    "issue_action",
    "reduction_summary",
    "unknown_table",
]

MAJOR_SECTION_HEADINGS = [
    "排气筒",
    "储罐",
    "装卸",
    "废水液面",
    "工艺过程无组织",
    "开停工检维修",
    "数字化",
    "整改建议",
]

OVERVIEW_HEADINGS = {
    "排气筒": ["排气筒概况"],
    "储罐": ["储罐分类统计"],
    "装卸": ["装卸设施统计情况"],
    "废水液面": ["废水液面统计情况"],
    "工艺过程无组织": ["工艺过程无组织统计"],
}

SUBSECTION_HEADINGS = {
    "排气筒概况",
    "储罐分类统计",
    "装卸设施统计情况",
    "废水液面统计情况",
    "工艺过程无组织统计",
    "问题识别",
    "减排空间识别",
}

ENTERPRISE_SECTION_HEADINGS = {
    "企业概述",
    "生产简介",
    "问题和减排空间",
    "重点异味点",
    "分项分析",
    "附件",
}


@dataclass
class ParseIssue:
    level: str
    code: str
    enterprise_name: str
    source_file: str
    message: str

    def as_dict(self) -> dict:
        return {
            "level": self.level,
            "code": self.code,
            "enterprise_name": self.enterprise_name,
            "source_file": self.source_file,
            "message": self.message,
        }


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").replace("\u3000", " ")).strip()


def normalize_enterprise_name(value: str, source_file: str = "") -> str:
    text = normalize_text(value)
    if not text and source_file:
        text = Path(source_file).stem
    text = re.sub(r"^\d+", "", text)
    text = re.sub(r"(总结材料|总计材料|总结报告|材料)$", "", text).strip()
    return text or normalize_text(value) or source_file


def natural_key(path: Path):
    """Sort: files with leading digit first, numeric ascending."""
    name = path.name
    m = re.match(r"^(\d+)", name)
    if m:
        return (0, int(m.group(1)), name)
    return (1, 999999, name)


def detect_input_dir(given: str | None) -> Path:
    if given:
        p = Path(given)
        if p.exists() and p.is_dir():
            return p
        raise FileNotFoundError(f"输入目录不存在: {given}")

    # Candidate paths — adjust for your environment
    candidates = [
        Path("/Volumes/a盘/project/排口总结材料0306-含炼化(1)"),
        Path("/Volumes/a盘/project/排口总结材料0306-含炼化"),
        Path("D:/project/排口总结材料0306-含炼化(1)"),
        Path("D:/project/排口总结材料0306-含炼化"),
    ]
    for c in candidates:
        if c.exists() and c.is_dir():
            return c

    # Last resort: glob
    for base in [Path("/Volumes/a盘/project"), Path("D:/project")]:
        if base.exists():
            for p in base.glob("*0306*"):
                if p.is_dir():
                    return p

    raise FileNotFoundError(
        "未自动发现输入目录，请使用 --input-dir 参数指定\n"
        "示例: python scripts/extract_outlet_structured.py \\\n"
        "  --input-dir '/Volumes/a盘/project/排口总结材料0306-含炼化(1)' \\\n"
        "  --output-dir data/outlet_structured_trial --limit 5"
    )


def extract_paragraphs(doc: Document) -> list[str]:
    return [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]


def table_preview(table: Table, max_rows: int = 3) -> str:
    rows = []
    for row in table.rows[:max_rows]:
        cells = [normalize_text(cell.text) for cell in row.cells]
        if any(cells):
            rows.append("|".join(cells))
    return " / ".join(rows)


def iter_document_blocks(doc: Document) -> list[dict]:
    blocks = []
    table_index = 0
    paragraph_index = 0
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            text = normalize_text(paragraph.text)
            if not text:
                continue
            paragraph_index += 1
            blocks.append(
                {
                    "block_type": "paragraph",
                    "block_index": len(blocks) + 1,
                    "paragraph_index": paragraph_index,
                    "table_index": "",
                    "text": text,
                    "style_name": paragraph.style.name if paragraph.style else "",
                }
            )
        elif isinstance(child, CT_Tbl):
            table_index += 1
            table = Table(child, doc)
            blocks.append(
                {
                    "block_type": "table",
                    "block_index": len(blocks) + 1,
                    "paragraph_index": "",
                    "table_index": table_index,
                    "text": table_preview(table),
                    "style_name": "",
                }
            )
    return blocks


def heading_key(text: str) -> str:
    return normalize_text(text).rstrip("：:")


def detect_heading_level(text: str, style_name: str, current_level: int = 1) -> tuple[int | None, str]:
    key = heading_key(text)
    style = style_name or ""

    if style == "Title":
        return 0, "style_title"

    m = re.match(r"Heading\s+(\d+)", style)
    if m:
        return int(m.group(1)), "style_heading"

    if style == "Caption":
        return min(current_level + 1, 6), "style_caption"

    if key in {"企业概述", "分项分析", "整改建议", "附件"}:
        return 1, "known_heading"
    if key in {"生产简介", "问题和减排空间", "重点异味点"} or key in MAJOR_SECTION_HEADINGS:
        return 2, "known_heading"
    if key in SUBSECTION_HEADINGS or key.startswith("附表") or key.startswith("附图"):
        return 3 if key in SUBSECTION_HEADINGS else 2, "known_heading"

    # Some source docs use plain Normal style for standalone sub-titles.
    if len(key) <= 40 and re.match(r"^\d+[、.．][^。；;]+$", key):
        return current_level if current_level >= 4 else min(current_level + 1, 6), "numbered_heading"
    if len(key) <= 40 and key.endswith(("：", ":")) and "。" not in key:
        return current_level if current_level >= 3 else min(current_level + 1, 6), "colon_heading"

    return None, ""


def blocks_to_text(blocks: list[dict], include_tables: bool = True) -> str:
    lines = []
    for block in blocks:
        if block["block_type"] == "paragraph":
            lines.append(block["text"])
        elif include_tables:
            table_text = f"[表{block['table_index']}]"
            if block.get("text"):
                table_text = f"{table_text} {block['text']}"
            lines.append(table_text)
    return "\n".join(lines).strip()


def table_refs(blocks: list[dict]) -> str:
    refs = []
    for block in blocks:
        if block["block_type"] == "table":
            refs.append(f"表{block['table_index']}")
    return "；".join(refs)


def extract_document_sections(doc: Document, enterprise_name: str, source_file: str) -> list[dict]:
    blocks = iter_document_blocks(doc)
    heading_positions = []
    current_level = 1
    for idx, block in enumerate(blocks):
        if block["block_type"] != "paragraph":
            continue
        level, method = detect_heading_level(block["text"], block.get("style_name", ""), current_level)
        if level is None:
            continue
        heading_positions.append((idx, level, method))
        current_level = level

    rows = []
    stack: list[tuple[int, str]] = []
    for order, (block_pos, level, method) in enumerate(heading_positions, start=1):
        block = blocks[block_pos]
        next_any = heading_positions[order][0] if order < len(heading_positions) else len(blocks)
        next_same_or_parent = len(blocks)
        for later_pos, later_level, _later_method in heading_positions[order:]:
            if later_level <= level:
                next_same_or_parent = later_pos
                break

        while stack and stack[-1][0] >= level:
            stack.pop()
        parent_path = " > ".join([title for _lvl, title in stack] + [block["text"]])
        stack.append((level, block["text"]))

        direct_blocks = blocks[block_pos + 1 : next_any]
        section_blocks = blocks[block_pos + 1 : next_same_or_parent]
        child_heading_positions = {
            pos: lvl for pos, lvl, _method in heading_positions if block_pos < pos < next_same_or_parent and lvl > level
        }
        child_headings = [blocks[pos]["text"] for pos in sorted(child_heading_positions)]

        rows.append(
            {
                "source_file": source_file,
                "table_index": "",
                "row_index": "",
                "enterprise_name": enterprise_name,
                "record_type": "document_section",
                "raw_fields": blocks_to_text(section_blocks),
                "heading_order": order,
                "block_index": block["block_index"],
                "paragraph_index": block["paragraph_index"],
                "heading_level": level,
                "heading_text": block["text"],
                "heading_style": block.get("style_name", ""),
                "heading_detection": method,
                "parent_path": parent_path,
                "direct_content": blocks_to_text(direct_blocks),
                "section_content": blocks_to_text(section_blocks),
                "child_headings": "；".join(child_headings),
                "direct_table_refs": table_refs(direct_blocks),
                "section_table_refs": table_refs(section_blocks),
            }
        )
    return rows


def is_attachment_heading(text: str) -> bool:
    key = heading_key(text)
    return key == "附件" or key.startswith("附表") or key.startswith("附图")


def first_block_after(
    paragraphs: list[str], heading: str, stop_headings: set[str]
) -> str:
    """Return text after the first matching heading until a stop heading."""
    start = None
    for idx, text in enumerate(paragraphs):
        if heading_key(text) == heading:
            start = idx + 1
            break
    if start is None:
        return ""

    block = []
    for text in paragraphs[start:]:
        key = heading_key(text)
        if key in stop_headings or is_attachment_heading(text):
            break
        block.append(text)
    return "\n".join(block).strip()


def extract_major_sections(paragraphs: list[str], enterprise_name: str, source_file: str) -> list[dict]:
    section_starts = []
    major_set = set(MAJOR_SECTION_HEADINGS)
    for idx, text in enumerate(paragraphs):
        key = heading_key(text)
        if key in major_set:
            section_starts.append((idx, key))

    sections = []
    for order, (start_idx, section_name) in enumerate(section_starts, start=1):
        end_idx = len(paragraphs)
        for next_idx, _next_name in section_starts[order:]:
            if next_idx > start_idx:
                end_idx = next_idx
                break
        body = []
        for text in paragraphs[start_idx + 1 : end_idx]:
            if is_attachment_heading(text):
                break
            body.append(text)

        overview_heading, overview_text = extract_overview(section_name, body)
        issue_text = extract_named_subsection(body, "问题识别")
        reduction_text = extract_named_subsection(body, "减排空间识别")
        if section_name in {"开停工检维修", "数字化", "整改建议"} and not overview_text:
            overview_text = "\n".join(t for t in body if heading_key(t) not in SUBSECTION_HEADINGS).strip()

        sections.append(
            {
                "source_file": source_file,
                "table_index": "",
                "row_index": "",
                "enterprise_name": enterprise_name,
                "record_type": "section_summary",
                "raw_fields": "\n".join(body),
                "section_order": order,
                "section_name": section_name,
                "overview_heading": overview_heading,
                "overview_text": overview_text,
                "issue_identification": issue_text,
                "reduction_space_identification": reduction_text,
            }
        )
    return sections


def extract_overview(section_name: str, body: list[str]) -> tuple[str, str]:
    expected = set(OVERVIEW_HEADINGS.get(section_name, []))
    if expected:
        for idx, text in enumerate(body):
            key = heading_key(text)
            if key in expected:
                return key, collect_until_subheading(body[idx + 1 :])

    leading = []
    for text in body:
        if heading_key(text) in SUBSECTION_HEADINGS:
            break
        leading.append(text)
    return "", "\n".join(leading).strip()


def extract_named_subsection(body: list[str], heading: str) -> str:
    for idx, text in enumerate(body):
        if heading_key(text) == heading:
            return collect_until_subheading(body[idx + 1 :])
    return ""


def collect_until_subheading(texts: list[str]) -> str:
    block = []
    for text in texts:
        if heading_key(text) in SUBSECTION_HEADINGS:
            break
        block.append(text)
    return "\n".join(block).strip()


def split_items(text: str) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []

    protected = text
    parts = [p.strip(" ，,；;。") for p in re.split(r"[；;]\s*", protected) if p.strip(" ，,；;。")]
    if len(parts) <= 1:
        numbered = re.split(r"(?=(?:\d+[、.．]|[一二三四五六七八九十]+[是、.．]))", protected)
        parts = [p.strip(" ，,；;。") for p in numbered if p.strip(" ，,；;。")]
    if len(parts) <= 1 and "。 " in protected:
        parts = [p.strip(" ，,；;。") for p in protected.split("。 ") if p.strip(" ，,；;。")]

    return parts or [text]


def make_section_items(sections: list[dict], field_name: str, record_type: str) -> list[dict]:
    rows = []
    for section in sections:
        for item_order, item_text in enumerate(split_items(section.get(field_name, "")), start=1):
            rows.append(
                {
                    "source_file": section["source_file"],
                    "table_index": "",
                    "row_index": "",
                    "enterprise_name": section["enterprise_name"],
                    "record_type": record_type,
                    "raw_fields": section.get(field_name, ""),
                    "section_order": section["section_order"],
                    "section_name": section["section_name"],
                    "item_order": item_order,
                    "item_text": item_text,
                    "is_no_issue_statement": "是" if re.search(r"无|未发现|无明显", item_text) else "",
                }
            )
    return rows


def parse_first_count(text: str, patterns: list[str]) -> str:
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            return m.group(1)
    return ""


def parse_outlet_overview(section: dict) -> dict:
    overview = section.get("overview_text", "")
    no_outlet = bool(
        re.search(r"(不涉及|无)[^。；]*(排气筒|有组织废气排放口|有组织排放口|废气排口)", overview)
    )
    total = parse_first_count(
        overview,
        [
            r"共有(?:排气筒|有组织排放口|排放口|排口)\s*(\d+)\s*[根个]",
            r"共有\s*(\d+)\s*[根个](?:排气筒|有组织排放口|排放口|排口)",
            r"目前有\s*(\d+)\s*个(?:废气)?排放口",
            r"(?:排气筒|有组织排放口|排放口|排口)\s*(\d+)\s*[根个]",
        ],
    )
    elevated = parse_first_count(
        overview,
        [
            r"高架排气筒[（(]?\s*50\s*米及以上\s*[）)]?\s*(\d+)\s*[根个]",
            r"50\s*米及以上[^。；，,]*?(\d+)\s*[根个]",
        ],
    )
    if no_outlet:
        total = total or "0"
        elevated = elevated or "0"

    composition_parts = re.findall(r"其中[^。；;]*?\d+\s*[根个][^。；;]*", overview)
    for clause in re.split(r"[。；;，,]\s*", overview):
        if re.search(r"(?:\d+\s*[根个]|[根个]\s*\d+)", clause) and re.search(
            r"排气筒|排放口|排口|火炬", clause
        ):
            composition_parts.append(clause)
    composition_parts = list(dict.fromkeys(p.strip() for p in composition_parts if p.strip()))
    dcs_status = ""
    if "DCS" in overview.upper():
        dcs_status = "无DCS" if re.search(r"无\s*DCS|没有\s*DCS", overview, flags=re.I) else "涉及DCS"

    return {
        "source_file": section["source_file"],
        "table_index": "",
        "row_index": "",
        "enterprise_name": section["enterprise_name"],
        "record_type": "outlet_overview",
        "raw_fields": overview,
        "section_order": section["section_order"],
        "total_outlet_count": total,
        "elevated_stack_50m_plus_count": elevated,
        "outlet_composition_text": "；".join(composition_parts),
        "dcs_status": dcs_status,
        "overview_text": overview,
        "issue_identification": section.get("issue_identification", ""),
        "reduction_space_identification": section.get("reduction_space_identification", ""),
    }


def extract_profile(paragraphs: list[str], source_file: str) -> dict:
    enterprise_name = normalize_enterprise_name(paragraphs[0] if paragraphs else "", source_file)
    rating = ""
    products_scale = ""
    reduction_text = ""

    for text in paragraphs:
        t = text
        if not rating and ("绩效评级" in t or "评定" in t):
            rating = t
        if not products_scale and ("产品" in t and "规模" in t):
            products_scale = t
        if not reduction_text and ("减排空间" in t or "减排量" in t or "减排" in t):
            reduction_text = t

    stop_after_intro = {"问题和减排空间", "重点异味点", "分项分析", "排气筒"}
    stop_after_problem = {"重点异味点", "分项分析", "排气筒"}
    stop_after_odor = {"分项分析", "排气筒"}

    return {
        "enterprise_name": enterprise_name,
        "source_file": source_file,
        "table_index": "",
        "row_index": "",
        "record_type": "enterprise_profile",
        "raw_fields": "\n".join(paragraphs[:40]),
        "production_intro": first_block_after(paragraphs, "生产简介", stop_after_intro),
        "problems_and_reduction_space": first_block_after(paragraphs, "问题和减排空间", stop_after_problem),
        "key_odor_points": first_block_after(paragraphs, "重点异味点", stop_after_odor),
        "rating_excerpt": rating,
        "products_scale_excerpt": products_scale,
        "reduction_excerpt": reduction_text,
    }


def make_unique_headers(headers: list[str]) -> list[str]:
    """Append __N suffix when duplicate header names appear."""
    result = []
    seen: dict[str, int] = defaultdict(int)
    for idx, raw in enumerate(headers, start=1):
        key = normalize_text(raw) or f"col_{idx}"
        seen[key] += 1
        if seen[key] > 1:
            key = f"{key}__{seen[key]}"
        result.append(key)
    return result


def classify_table(headers: list[str]) -> str:
    """
    Classify table type from its header row.
    Headers are already normalized via make_unique_headers.
    We search in the raw (non-normalized) headers joined string to avoid
    issues with duplicate-header deduplication.
    """
    h = "|".join(headers)

    # ── issue_action ──────────────────────────────────────────────
    # 必有：问题内容描述 + (建议措施 | 企业已采取措施)
    if "问题内容描述" in h:
        if "建议措施" in h or "企业已采取措施" in h:
            return "issue_action"

    # ── reduction_summary ──────────────────────────────────────────
    # 必有：源项 + 整改建议 + 减排量/减排
    if "源项" in h and "整改建议" in h:
        if "减排" in h:
            return "reduction_summary"
        # Fallback: just source items + suggestions (some files omit "减排" in header)
        return "reduction_summary"

    # ── fugitive_source ────────────────────────────────────────────
    if "无组织废气名称" in h:
        return "fugitive_source"

    # ── organized_outlet ───────────────────────────────────────────
    # 有组织排放口名称 / 排气筒名称 + 污染物/检测因子
    if "有组织排放口名称" in h:
        return "organized_outlet"
    if "排气筒名称" in h and "污染物" in h:
        return "organized_outlet"
    if "排气筒名称" in h and "检测因子" in h:
        return "organized_outlet"
    if "排放口编号" in h and "主要污染物名称" in h:
        return "organized_outlet"
    if "排放口编号" in h and "检测因子" in h:
        return "organized_outlet"

    # ── tank ───────────────────────────────────────────────────────
    # 储罐名称（含储罐位号）/ 罐型 + 周转量/容积
    if "储罐" in h and "罐型" in h:
        return "tank"
    if "储罐" in h and "周转量" in h:
        return "tank"

    # ── loading ────────────────────────────────────────────────────
    # 装卸设施名称 + (周转量 | 装载量 | 装卸量)
    if "装卸设施名称" in h or "装卸" in h:
        if "周转量" in h or "装载量" in h or "装卸量" in h:
            return "loading"
        if "装卸方式" in h:
            return "loading"

    # ── wastewater_surface ─────────────────────────────────────────
    # 废水液面 / 排放源名称(含"液面"/"敞开") / 装置名称+设施名称+涉VOCs
    if "废水" in h and "液面" in h:
        return "wastewater_surface"
    if "排放源名称" in h and "所在位置" in h:
        return "wastewater_surface"
    if "排放源名称" in h and "是否加盖" in h:
        return "wastewater_surface"
    if "装置名称" in h and "设施名称" in h and "涉VOCs" in h:
        return "wastewater_surface"

    return "unknown_table"


def get_field_by_keywords(row: dict, keywords: list[str]) -> str:
    for key, value in row.items():
        if key in {"source_file", "table_index", "row_index", "enterprise_name", "record_type", "raw_fields"}:
            continue
        clean_key = re.sub(r"__\d+$", "", str(key))
        if all(word in clean_key for word in keywords):
            return value
    return ""


def first_matching_field(row: dict, keyword_groups: list[list[str]]) -> str:
    for keywords in keyword_groups:
        value = get_field_by_keywords(row, keywords)
        if value:
            return value
    return ""


def add_standard_fields(row: dict, table_type: str) -> None:
    """Add common analysis fields while keeping original table columns intact."""
    if table_type == "organized_outlet":
        row["facility_name_raw"] = first_matching_field(
            row, [["有组织排放口名称"], ["排气筒名称"], ["排放口名称"]]
        )
        row["facility_code"] = first_matching_field(row, [["排放口编号"], ["编号"]])
        row["pollutant_category"] = first_matching_field(
            row, [["主要污染物名称"], ["污染物"], ["检测因子"]]
        )
        row["treatment_process"] = first_matching_field(row, [["废气处理工艺"], ["治理工艺"], ["废气去向"]])
        row["monitoring_method"] = first_matching_field(row, [["是否设置在线"], ["在线"], ["监测"]])
    elif table_type == "tank":
        row["facility_name_raw"] = first_matching_field(row, [["储罐名称"], ["储罐位号"], ["储罐"]])
        row["facility_code"] = first_matching_field(row, [["编号"], ["位号"]])
        row["pollutant_category"] = first_matching_field(row, [["物料"], ["介质"]])
        row["treatment_process"] = first_matching_field(row, [["废气去向"], ["处理"], ["治理"]])
    elif table_type == "loading":
        row["facility_name_raw"] = first_matching_field(row, [["装卸设施名称"], ["装卸"]])
        row["process_stage"] = first_matching_field(row, [["装卸方式"], ["工段"], ["环节"]])
        row["pollutant_category"] = first_matching_field(row, [["物料"], ["介质"]])
        row["treatment_process"] = first_matching_field(row, [["废气去向"], ["处理"], ["治理"]])
    elif table_type == "wastewater_surface":
        row["facility_name_raw"] = first_matching_field(row, [["排放源名称"], ["设施名称"], ["液面"]])
        row["process_stage"] = first_matching_field(row, [["所在位置"], ["装置名称"]])
        row["monitoring_method"] = first_matching_field(row, [["是否有明显异味"], ["VOCs检测"], ["检测浓度"]])
        row["treatment_process"] = first_matching_field(row, [["废气是否收集处理"], ["是否加盖密闭"]])
    elif table_type == "fugitive_source":
        row["facility_name_raw"] = first_matching_field(row, [["无组织废气名称"], ["设备名称"]])
        row["process_stage"] = first_matching_field(row, [["工段或工艺名称"], ["工段"], ["工艺"]])
        row["monitoring_method"] = first_matching_field(row, [["预计收集率"], ["废气是否收集"]])
        row["treatment_process"] = first_matching_field(row, [["废气去向"], ["废气收集方式"]])
    elif table_type == "issue_action":
        row["facility_name_raw"] = first_matching_field(row, [["问题点位"], ["源项"], ["位置"]])
        row["process_stage"] = first_matching_field(row, [["分类"], ["环节"]])
        row["treatment_process"] = first_matching_field(row, [["建议措施"], ["企业已采取措施"], ["整改"]])
    elif table_type == "reduction_summary":
        row["facility_name_raw"] = first_matching_field(row, [["源项"], ["设施"], ["排口"]])
        row["treatment_process"] = first_matching_field(row, [["整改建议"], ["措施"]])
        row["pollutant_category"] = first_matching_field(row, [["污染物"], ["减排量"]])


def is_repeat_header(row: list[str], headers: list[str]) -> bool:
    """Detect sub-header / year-split rows that duplicate the header content."""
    row_text = normalize_text(" ".join(row))
    header_text = normalize_text(" ".join(headers))
    if not row_text or not header_text:
        return False
    header_words = set(w for w in re.split(r"[\s|,，、]+", header_text) if len(w) > 1)
    if not header_words:
        return False
    match_count = sum(1 for w in header_words if w in row_text)
    return match_count / len(header_words) >= 0.7


def parse_table(
    table, source_file: str, table_index: int, enterprise_name: str
) -> tuple[str, list[dict]]:
    """Extract rows from a docx table, classify type, return (type, rows)."""
    raw_rows: list[list[str]] = []
    for row in table.rows:
        cells = [normalize_text(cell.text) for cell in row.cells]
        deduped = []
        prev = None
        for c in cells:
            if c != prev:
                deduped.append(c)
            prev = c
        raw_rows.append(deduped)

    if len(raw_rows) < 2:
        return ("unknown_table", [])

    header_row = raw_rows[0]
    if all(c == "" for c in header_row) and len(raw_rows) > 2:
        header_row = raw_rows[1]
        data_rows = raw_rows[2:]
    else:
        data_rows = raw_rows[1:]

    headers = make_unique_headers(header_row)
    table_type = classify_table(header_row)

    records = []
    first_data_row_index = 3 if all(c == "" for c in raw_rows[0]) and len(raw_rows) > 2 else 2
    for ri, raw in enumerate(data_rows, start=first_data_row_index):
        if not any(c.strip() for c in raw):
            continue
        if is_repeat_header(raw, header_row):
            continue

        row_dict = {"source_file": source_file, "table_index": table_index, "row_index": ri}
        for ci, h in enumerate(headers):
            val = raw[ci].strip() if ci < len(raw) else ""
            row_dict[h] = val
        row_dict["enterprise_name"] = enterprise_name
        row_dict["record_type"] = table_type
        row_dict["raw_fields"] = "|".join(raw)
        add_standard_fields(row_dict, table_type)
        records.append(row_dict)

    return (table_type, records)


def process_file(docx_path: Path, issues: list[ParseIssue]) -> dict:
    """Process one .docx file and return structured records by type."""
    source_file = docx_path.name
    result: dict[str, list[dict]] = defaultdict(list)

    try:
        doc = Document(str(docx_path))
    except Exception as e:
        issues.append(ParseIssue("ERROR", "DOCX_OPEN", "", source_file, str(e)))
        return result

    paragraphs = extract_paragraphs(doc)
    enterprise_name = normalize_enterprise_name(paragraphs[0] if paragraphs else "", source_file)
    result["document_section"].extend(extract_document_sections(doc, enterprise_name, source_file))
    profile = extract_profile(paragraphs, source_file)
    result["enterprise_profile"].append(profile)
    sections = extract_major_sections(paragraphs, enterprise_name, source_file)
    result["section_summary"].extend(sections)
    result["issue_identification"].extend(
        make_section_items(sections, "issue_identification", "issue_identification")
    )
    result["reduction_space_identification"].extend(
        make_section_items(sections, "reduction_space_identification", "reduction_space_identification")
    )
    for section in sections:
        if section.get("section_name") == "排气筒":
            result["outlet_overview"].append(parse_outlet_overview(section))

    for ti, table in enumerate(doc.tables, start=1):
        table_type, rows = parse_table(table, source_file, ti, enterprise_name)
        result[table_type].extend(rows)
        if table_type == "unknown_table" and rows:
            issues.append(
                ParseIssue(
                    "WARN",
                    "UNKNOWN_TABLE",
                    enterprise_name,
                    source_file,
                    f"Table {ti}: {rows[0].get('raw_fields', '')[:120]}",
                )
            )

    return result


def ordered_fields(rows: list[dict]) -> list[str]:
    field_set: set[str] = set()
    first_seen = []
    for r in rows:
        for key in r.keys():
            if key not in field_set:
                field_set.add(key)
                first_seen.append(key)
    standard = [
        "source_file",
        "table_index",
        "row_index",
        "enterprise_name",
        "record_type",
        "heading_order",
        "heading_level",
        "heading_text",
        "parent_path",
        "section_order",
        "section_name",
        "item_order",
        "direct_content",
        "section_content",
        "child_headings",
        "direct_table_refs",
        "section_table_refs",
        "facility_name_raw",
        "facility_code",
        "process_stage",
        "pollutant_category",
        "monitoring_method",
        "treatment_process",
        "raw_fields",
    ]
    seen = set()
    final_fields = []
    for f in standard + first_seen:
        if f in field_set and f not in seen:
            seen.add(f)
            final_fields.append(f)
    return final_fields


def write_csvs(all_records: dict[str, list[dict]], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    for table_type in TYPE_ORDER:
        rows = all_records.get(table_type, [])
        csv_path = output_dir / f"{table_type}.csv"
        if not rows:
            csv_path.write_text("", encoding="utf-8-sig")
            continue

        final_fields = ordered_fields(rows)

        with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=final_fields, extrasaction="ignore")
            writer.writeheader()
            writer.writerows(rows)

        print(f"  {table_type}: {len(rows)} rows -> {csv_path.name}")


def write_workbook(all_records: dict[str, list[dict]], output_dir: Path) -> None:
    try:
        from openpyxl import Workbook
    except ImportError:
        print("  workbook skipped: openpyxl not installed")
        return

    wb = Workbook()
    default = wb.active
    wb.remove(default)

    for table_type in TYPE_ORDER:
        ws = wb.create_sheet(title=table_type[:31])
        rows = all_records.get(table_type, [])
        if not rows:
            continue
        fields = ordered_fields(rows)
        ws.append(fields)
        for row in rows:
            ws.append([row.get(field, "") for field in fields])
        ws.freeze_panes = "A2"

    xlsx_path = output_dir / "outlet_structured_summary.xlsx"
    wb.save(xlsx_path)
    print(f"  workbook: {xlsx_path.name}")


def write_run_summary(
    all_records: dict[str, list[dict]], issues: list[ParseIssue], output_dir: Path, input_dir: Path, file_count: int
) -> None:
    total_data = sum(len(all_records.get(t, [])) for t in TYPE_ORDER if t != "enterprise_profile")
    unknown = len(all_records.get("unknown_table", []))
    summary = {
        "input_dir": str(input_dir),
        "output_dir": str(output_dir),
        "file_count": file_count,
        "tables": {t: len(all_records.get(t, [])) for t in TYPE_ORDER},
        "total_data_rows": total_data,
        "unknown_table_rate": round(unknown / total_data, 4) if total_data else 0,
        "issue_count": len(issues),
    }
    (output_dir / "run_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print("  run_summary: run_summary.json")


def write_issues(issues: list[ParseIssue], output_dir: Path) -> None:
    csv_path = output_dir / "qa_issues.csv"
    if not issues:
        csv_path.write_text("", encoding="utf-8-sig")
        return
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=["level", "code", "enterprise_name", "source_file", "message"])
        writer.writeheader()
        for issue in issues:
            writer.writerow(issue.as_dict())
    print(f"  qa_issues: {len(issues)} issues")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract structured outlet data from .docx files")
    parser.add_argument("--input-dir", default=None, help="Directory containing .docx files")
    parser.add_argument("--output-dir", default="data/outlet_structured", help="Output directory for CSVs")
    parser.add_argument("--limit", type=int, default=0, help="Only process first N files (0=all)")
    parser.add_argument("--all", action="store_true", help="Process all files (alias for --limit 0)")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.all:
        args.limit = 0

    input_dir = detect_input_dir(args.input_dir)
    output_dir = Path(args.output_dir)

    docx_files = sorted(input_dir.glob("*.docx"), key=natural_key)
    if args.limit > 0:
        docx_files = docx_files[: args.limit]

    print(f"Input dir:  {input_dir}")
    print(f"Output dir: {output_dir}")
    print(f"Files to process: {len(docx_files)}")

    all_records: dict[str, list[dict]] = defaultdict(list)
    issues: list[ParseIssue] = []

    for i, fp in enumerate(docx_files, 1):
        print(f"\n[{i}/{len(docx_files)}] {fp.name}")
        records = process_file(fp, issues)
        for k, v in records.items():
            all_records[k].extend(v)

    print(f"\n{'='*60}")
    print("Writing outputs...")
    write_csvs(all_records, output_dir)
    write_issues(issues, output_dir)
    write_workbook(all_records, output_dir)
    write_run_summary(all_records, issues, output_dir, input_dir, len(docx_files))

    print(f"\n{'='*60}")
    print("Summary:")
    for t in TYPE_ORDER:
        n = len(all_records.get(t, []))
        if n > 0:
            print(f"  {t}: {n}")
    total_data = sum(len(all_records.get(t, [])) for t in TYPE_ORDER if t != "enterprise_profile")
    total_enterprises = len(all_records.get("enterprise_profile", []))
    print(f"  Total enterprises: {total_enterprises}")
    print(f"  Total data rows: {total_data}")
    unknown = len(all_records.get("unknown_table", []))
    if total_data > 0:
        print(f"  Unknown table rate: {unknown}/{total_data} = {unknown/total_data:.1%}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
